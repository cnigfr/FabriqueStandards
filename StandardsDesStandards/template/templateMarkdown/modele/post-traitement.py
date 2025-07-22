import sys
import os
import zipfile
import shutil
import tempfile
from docx import Document
from docx.shared import Pt
from lxml import etree
import xml.etree.ElementTree as ET

# Namespace XML Word
NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
ET.register_namespace('w', NS['w'])

def set_table_font_size(docx_path, font_size=9):
    doc = Document(docx_path)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)
    try:
        doc.save(docx_path)
        print(f"✅ Police des tableaux mise à {font_size} pt.")
    except PermissionError:
        print(f"❌ Impossible d'enregistrer '{docx_path}'. Ferme le fichier s'il est ouvert dans Word.")
        sys.exit(1)

def remove_before_fiche(xml_data):
    root = ET.fromstring(xml_data)
    body = root.find('w:body', NS)

    new_body = []
    found = False

    for elem in list(body):
        texts = elem.findall(".//w:t", NS)
        full_text = " ".join(t.text for t in texts if t.text).lower()
        if not found and "fiche descriptive" in full_text:
            found = True
            new_body.append(elem)
        elif found:
            new_body.append(elem)

    body.clear()
    for elem in new_body:
        body.append(elem)

    return ET.tostring(root, encoding='utf-8', xml_declaration=True)

def enable_autofit_tables(xml_data):
    tree = etree.fromstring(xml_data)

    for tbl in tree.xpath('//w:tbl', namespaces=NS):
        tblW = tbl.find('w:tblPr/w:tblW', namespaces=NS)
        if tblW is not None:
            tbl.find('w:tblPr', namespaces=NS).remove(tblW)

        tblLayout = tbl.find('w:tblPr/w:tblLayout', namespaces=NS)
        if tblLayout is not None:
            tbl.find('w:tblPr', namespaces=NS).remove(tblLayout)

        tblPr = tbl.find('w:tblPr', namespaces=NS)
        layout = etree.Element('{%s}tblLayout' % NS['w'])
        layout.set('{%s}type' % NS['w'], 'autofit')
        tblPr.append(layout)

    return etree.tostring(tree, xml_declaration=True, encoding='utf-8')

def postprocess_docx(input_path):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        temp_path = tmp.name

    with zipfile.ZipFile(input_path, 'r') as zin:
        with zipfile.ZipFile(temp_path, 'w') as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    data = remove_before_fiche(data)
                    data = enable_autofit_tables(data)
                zout.writestr(item, data)

    shutil.move(temp_path, input_path)
    print("✅ Post-traitement XML terminé.")

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("❌ Utilisation : python a.py Document.docx")
        sys.exit(1)

    fichier = sys.argv[1]

    if not os.path.isfile(fichier):
        print(f"❌ Le fichier '{fichier}' n'existe pas.")
        sys.exit(1)

    # Étape 1 : modifier la taille de police dans les tableaux
    set_table_font_size(fichier)

    # Étape 2 : traitement XML (nettoyage + autofit)
    postprocess_docx(fichier)
